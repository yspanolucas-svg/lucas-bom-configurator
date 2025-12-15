import io
import os
import re
import hashlib
from dataclasses import dataclass
from typing import List, Optional, Dict

import pandas as pd
import streamlit as st
from PyPDF2 import PdfReader

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm

from docx import Document  # python-docx
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# =========================================================
# DATA MODEL
# =========================================================
@dataclass
class PDFRow:
    key: str
    label: str
    value: str
    source: str
    section: str  # "main"


# =========================================================
# HELPERS
# =========================================================
def _hash_bytes(b: Optional[bytes]) -> str:
    if not b:
        return ""
    return hashlib.sha1(b).hexdigest()


def _norm_label(s: str) -> str:
    return re.sub(r"\s+", " ", str(s).strip().lower())


# =========================================================
# PRESET (from repo) - supports your "1 line = 1 label" format
# =========================================================
def load_fusion_preset_labels_from_repo(preset_path: str = "presets/fusion_default.csv") -> set[str]:
    """
    Supporte 2 formats :
    1) CSV 'propre' (avec séparateurs) contenant une colonne 'label'
    2) Fichier simple : 1 ligne = 1 label (format actuel)
       Exemple:
         label
         Maximum speed on the y-axis
         ...
    """
    if not os.path.exists(preset_path):
        return set()

    # 1) tentative CSV classique
    try:
        df = pd.read_csv(preset_path)
        # Si c'est un vrai CSV avec plusieurs colonnes (donc séparateurs présents)
        if "label" in df.columns and len(df.columns) > 1:
            return set(df["label"].dropna().apply(_norm_label))
    except Exception:
        pass

    # 2) fallback robuste : 1 ligne = 1 label
    labels = set()
    with open(preset_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.lower() == "label":
                continue
            labels.add(_norm_label(line))

    return labels


def apply_fusion_presets_from_repo(df: pd.DataFrame, preset_labels: set[str]) -> pd.DataFrame:
    """
    Pre-check rows based on exact label match (normalized).
    Applies only Keep=True; doesn't touch Value.
    """
    if df is None or df.empty or not preset_labels:
        return df

    labels_norm = df["Label"].astype(str).apply(_norm_label)
    mask = labels_norm.isin(preset_labels)
    df.loc[mask, "Keep"] = True
    return df


# =========================================================
# PDF TEXT EXTRACTION
# =========================================================
def _pdf_text(pdf_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    chunks = []
    for p in reader.pages:
        chunks.append(p.extract_text() or "")
    txt = "\n".join(chunks)
    return txt.replace("\r", "\n")


def _parse_pairs_all(text: str) -> List[tuple]:
    """
    Parse des lignes de tableau.
    Gère :
      - "Libellé  Valeur" (double espaces)
      - "Robot weight 250.0 kg" (un seul espace)
      - fallback : dernier token = valeur
    """
    rows = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue

        # Priorité : "Label  Value"
        m = re.match(r"^(.*?)[ ]{2,}(.+)$", line)
        if m:
            rows.append((m.group(1).strip(), m.group(2).strip()))
            continue

        # Souvent : label + numérique à la fin
        m2 = re.match(r"^(.*?)([-+]?\d[\d.,]*\s*.*)$", line)
        if m2 and len(m2.group(1).strip()) >= 3:
            rows.append((m2.group(1).strip(), m2.group(2).strip()))
            continue

        # fallback
        parts = line.split(" ")
        if len(parts) >= 2:
            rows.append((" ".join(parts[:-1]).strip(), parts[-1].strip()))
        else:
            rows.append((line, ""))

    return rows


def extract_pdf_rows(pdf_bytes: Optional[bytes], doc_id: str, source_display: str) -> List[PDFRow]:
    """
    Keys stables basés sur doc_id + label + compteur.
    source_display = nom utilisateur (Ensemble 1 / Ensemble 2).
    """
    if not pdf_bytes:
        return []

    text = _pdf_text(pdf_bytes)
    pairs = _parse_pairs_all(text)

    rows: List[PDFRow] = []
    counter: Dict[str, int] = {}

    for label, value in pairs:
        label = (label or "").strip()
        value = (value or "").strip()
        if len(label) < 3:
            continue

        base = re.sub(r"\s+", "_", label.lower()).strip("_")
        counter[base] = counter.get(base, 0) + 1
        key = f"{doc_id}::{base}::{counter[base]}"

        rows.append(PDFRow(
            key=key,
            label=label,
            value=value,
            source=source_display,
            section="main",
        ))

    return rows


# =========================================================
# OUTPUT PDF (REPORTLAB) - bandeau Lucas + multi-pages lisibles
# =========================================================
def _draw_header_pdf(c: canvas.Canvas, title: str = "") -> float:
    w, h = A4

    # Bandeau noir
    c.setFillColorRGB(0.0, 0.0, 0.0)
    c.rect(0, h - 22 * mm, w, 22 * mm, stroke=0, fill=1)

    # Liseré rouge
    c.setFillColorRGB(0.85, 0.0, 0.0)
    c.rect(0, h - 22 * mm, w, 3.5 * mm, stroke=0, fill=1)

    # Texte blanc
    c.setFillColorRGB(1.0, 1.0, 1.0)
    c.setFont("Helvetica-Bold", 13)
    c.drawString(12 * mm, h - 14 * mm, "LUCAS ROBOTIC SYSTEM")

    c.setFont("Helvetica", 10)
    right = "FICHE TECHNIQUE"
    if title:
        right = f"{right}  –  {title}"
    c.drawRightString(w - 12 * mm, h - 14 * mm, right)

    return h - 28 * mm


def _wrap_simple(text: str, max_chars: int) -> List[str]:
    text = (text or "").strip()
    if not text:
        return [""]
    words = text.split()
    lines, cur = [], ""
    for w in words:
        if len(cur) + len(w) + 1 <= max_chars:
            cur = (cur + " " + w).strip()
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines or [""]


def _draw_table_header_pdf(c: canvas.Canvas, col_label: float, col_value: float, col_source: float, y: float) -> float:
    # Important : reset noir à chaque page
    c.setFillColorRGB(0.0, 0.0, 0.0)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(col_label, y, "Variable")
    c.drawString(col_value, y, "Valeur")
    c.drawString(col_source, y, "Source")
    y -= 4 * mm

    c.setStrokeColorRGB(0.7, 0.7, 0.7)
    c.setLineWidth(0.4)
    c.line(col_label, y, (A4[0] - 12 * mm), y)
    y -= 4 * mm

    c.setFont("Helvetica", 9)
    c.setFillColorRGB(0.0, 0.0, 0.0)
    return y


def _draw_table_pdf(c: canvas.Canvas, rows: List[PDFRow], y_start: float) -> float:
    x = 12 * mm
    w = A4[0] - 24 * mm
    y = y_start

    col_label = x
    col_value = x + w * 0.58
    col_source = x + w * 0.85

    y = _draw_table_header_pdf(c, col_label, col_value, col_source, y)

    for r in rows:
        lab_lines = _wrap_simple(r.label, 55)
        val_lines = _wrap_simple(r.value, 30)
        n = max(len(lab_lines), len(val_lines), 1)
        row_h = n * 4.2 * mm

        if y - row_h < 15 * mm:
            c.showPage()
            y = _draw_header_pdf(c, title="")
            y = _draw_table_header_pdf(c, col_label, col_value, col_source, y)

        c.setFillColorRGB(0.0, 0.0, 0.0)
        c.setFont("Helvetica", 9)

        for i in range(n):
            c.drawString(col_label, y - i * 4.2 * mm, lab_lines[i] if i < len(lab_lines) else "")
            c.drawString(col_value, y - i * 4.2 * mm, val_lines[i] if i < len(val_lines) else "")
            if i == 0:
                c.drawString(col_source, y, r.source)

        y -= row_h + 1.5 * mm

    return y


def build_output_pdf(selected_rows: List[PDFRow]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    y = _draw_header_pdf(c, title="")
    _draw_table_pdf(c, selected_rows, y)
    c.save()
    return buf.getvalue()


# =========================================================
# OUTPUT WORD (DOCX) - bandeau Lucas + éditable
# =========================================================
def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for k, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = OxmlElement(f"w:{k}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_lucas_banner_docx(doc: Document, right_title: str = "FICHE TECHNIQUE") -> None:
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in t.rows[0].cells:
        _set_cell_shading(cell, "000000")
        _set_cell_margins(cell, top=120, bottom=120, start=180, end=180)

    for cell in t.rows[1].cells:
        _set_cell_shading(cell, "D90000")
        _set_cell_margins(cell, top=20, bottom=20, start=0, end=0)

    t.rows[1].cells[0].merge(t.rows[1].cells[1])

    p_left = t.rows[0].cells[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_left.add_run("LUCAS ROBOTIC SYSTEM")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(13)

    p_right = t.rows[0].cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p_right.add_run(right_title)
    run2.bold = True
    run2.font.color.rgb = RGBColor(255, 255, 255)
    run2.font.size = Pt(11)

    doc.add_paragraph("")


def build_output_docx(selected_rows: List[PDFRow]) -> bytes:
    doc = Document()
    _add_lucas_banner_docx(doc, right_title="FICHE TECHNIQUE")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Variable"
    hdr[1].text = "Valeur"
    hdr[2].text = "Source"

    for r in selected_rows:
        row = table.add_row().cells
        row[0].text = str(r.label)
        row[1].text = str(r.value)
        row[2].text = str(r.source)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =========================================================
# STREAMLIT UI
# =========================================================
def render_pdf_lab_panel():
    st.subheader("PDF Lab – Nettoyage / Assemblage")

    with st.expander("Mode d’emploi (aperçu rapide)", expanded=True):
        st.markdown(
            """
**Objectif :** générer une fiche technique claire et cohérente à partir d’un ou deux PDF,
en supprimant les illustrations, la nomenclature (BOM) et les blocs non pertinents.

**Déroulé**
1) Choisir le mode : **Nettoyage (1 PDF)** ou **Assemblage (2 PDF)**  
2) Charger le ou les PDF  
3) Renommer les ensembles si nécessaire (ex : *Ensemble A / Ensemble B*)  
4) Sélectionner les caractéristiques à conserver  
5) Modifier si besoin certaines valeurs (course, vitesse, charge, etc.)  
6) Exporter la fiche finale en **PDF** et/ou **Word (DOCX)**

L’outil est conçu pour accélérer la préparation des offres et sécuriser le contenu
des fiches techniques utilisées par l’avant-projet et le commerce.
            """
        )

    mode = st.radio(
        "Mode",
        ["Nettoyage (1 PDF)", "Assemblage (2 PDF)"],
        horizontal=True,
        key="pdf_mode",
        help="Nettoyage = sélectionner des caractéristiques depuis un seul PDF. Assemblage = sélectionner des caractéristiques depuis deux PDF et produire une fiche unique.",
    )

    st.caption("Renommez les ensembles pour rester générique (orientation inconnue à ce stade).")
    coln1, coln2 = st.columns(2)
    with coln1:
        src1_name = st.text_input("Nom ensemble 1", value="Ensemble 1", key="src1_name")
    with coln2:
        src2_name = st.text_input("Nom ensemble 2", value="Ensemble 2", key="src2_name")

    if mode == "Nettoyage (1 PDF)":
        st.info("Chargez un PDF, sélectionnez les caractéristiques à conserver, puis exportez.")
        pdf1 = st.file_uploader("PDF – Ensemble 1", type=["pdf"], key="pdf1")
        pdf2 = None
    else:
        st.info("Chargez deux PDF, sélectionnez ce que vous souhaitez conserver dans chacun, puis exportez une fiche unique.")
        c1, c2 = st.columns(2)
        with c1:
            pdf1 = st.file_uploader("PDF – Ensemble 1", type=["pdf"], key="pdf1")
        with c2:
            pdf2 = st.file_uploader("PDF – Ensemble 2", type=["pdf"], key="pdf2")

    if not pdf1:
        st.warning("Veuillez charger au moins un PDF pour continuer.")
        return
    if mode == "Assemblage (2 PDF)" and not pdf2:
        st.warning("Veuillez charger le second PDF pour l’assemblage.")
        return

    b1 = pdf1.read()
    b2 = pdf2.read() if pdf2 else None

    h1 = _hash_bytes(b1)
    h2 = _hash_bytes(b2)
    doc1_id = f"doc1::{h1}"
    doc2_id = f"doc2::{h2}" if b2 else ""

    rows1_key = f"pdf_lab::rows::{doc1_id}"
    df1_key = f"pdf_lab::df::{doc1_id}"
    rows2_key = f"pdf_lab::rows::{doc2_id}"
    df2_key = f"pdf_lab::df::{doc2_id}"

    preset_labels = set()
    if mode == "Assemblage (2 PDF)":
        preset_labels = load_fusion_preset_labels_from_repo("presets/fusion_default.csv")

    # INIT DOC1 (apply preset here in Assemblage)
    if rows1_key not in st.session_state:
        rows1 = extract_pdf_rows(b1, doc_id=doc1_id, source_display=src1_name)
        st.session_state[rows1_key] = rows1

        df1_init = pd.DataFrame([{
            "Keep": False,
            "Label": r.label,
            "Value": r.value,
            "Source": r.source,
            "_key": r.key,
        } for r in rows1])

        if mode == "Assemblage (2 PDF)" and preset_labels:
            df1_init = apply_fusion_presets_from_repo(df1_init, preset_labels)

        st.session_state[df1_key] = df1_init

    # INIT DOC2 (apply preset here in Assemblage)
    if b2 and rows2_key not in st.session_state:
        rows2 = extract_pdf_rows(b2, doc_id=doc2_id, source_display=src2_name)
        st.session_state[rows2_key] = rows2

        df2_init = pd.DataFrame([{
            "Keep": False,
            "Label": r.label,
            "Value": r.value,
            "Source": r.source,
            "_key": r.key,
        } for r in rows2])

        if mode == "Assemblage (2 PDF)" and preset_labels:
            df2_init = apply_fusion_presets_from_repo(df2_init, preset_labels)

        st.session_state[df2_key] = df2_init

    df1 = st.session_state[df1_key]
    rows1 = st.session_state[rows1_key]
    if b2:
        df2 = st.session_state[df2_key]
        rows2 = st.session_state[rows2_key]
    else:
        df2 = pd.DataFrame(columns=df1.columns)
        rows2 = []

    # Rename sources (cosmetic)
    df1["Source"] = src1_name
    for r in rows1:
        r.source = src1_name
    if b2:
        df2["Source"] = src2_name
        for r in rows2:
            r.source = src2_name

    n_checked_1 = int(df1["Keep"].sum()) if "Keep" in df1.columns else 0
    n_checked_2 = int(df2["Keep"].sum()) if (b2 and "Keep" in df2.columns) else 0

    st.success(
        f"Caractéristiques détectées — {src1_name} : {len(df1)} (sélectionnées : {n_checked_1})"
        + (f" | {src2_name} : {len(df2)} (sélectionnées : {n_checked_2})" if b2 else "")
    )

    st.caption("Utilisez le champ de recherche pour retrouver rapidement une caractéristique (intitulé ou valeur).")
    search = st.text_input("Recherche (intitulé / valeur)", value="", key="pdf_search")

    tab_titles = [f"Ensemble 1 – {src1_name}"]
    if b2:
        tab_titles.append(f"Ensemble 2 – {src2_name}")
    tabs = st.tabs(tab_titles)

    def _apply_editor_changes(df: pd.DataFrame, edited: pd.DataFrame) -> pd.DataFrame:
        if edited is None or edited.empty:
            return df

        ed_map = {}
        for _, r in edited.iterrows():
            ed_map[str(r["_key"])] = {"Keep": bool(r["Keep"]), "Value": str(r["Value"])}

        for i in range(len(df)):
            k = str(df.at[i, "_key"])
            if k in ed_map:
                df.at[i, "Keep"] = ed_map[k]["Keep"]
                df.at[i, "Value"] = ed_map[k]["Value"]
        return df

    def editor(df: pd.DataFrame, key_prefix: str) -> pd.DataFrame:
        show_only_checked = st.toggle(
            "Afficher uniquement les lignes sélectionnées",
            value=False,
            key=f"{key_prefix}_show_checked",
        )

        df_view = df.copy()
        if search.strip():
            s = search.strip().lower()
            df_view = df_view[
                df_view["Label"].str.lower().str.contains(s, na=False)
                | df_view["Value"].astype(str).str.lower().str.contains(s, na=False)
            ]

        if show_only_checked:
            df_view = df_view[df_view["Keep"] == True]
            if df_view.empty:
                st.info("Aucune ligne sélectionnée dans cet ensemble pour le moment.")
                return df

        cA, cB, _ = st.columns([1, 1, 6])
        with cA:
            do_check_all = st.button(
                "Tout sélectionner (filtré)",
                key=f"{key_prefix}_all_on",
                help="Sélectionne toutes les lignes visibles (après recherche/filtre).",
            )
        with cB:
            do_uncheck_all = st.button(
                "Tout désélectionner (filtré)",
                key=f"{key_prefix}_all_off",
                help="Désélectionne toutes les lignes visibles (après recherche/filtre).",
            )

        if do_check_all:
            df.loc[df_view.index, "Keep"] = True
        if do_uncheck_all:
            df.loc[df_view.index, "Keep"] = False

        # rebuild view after button actions
        df_view = df.copy()
        if search.strip():
            s = search.strip().lower()
            df_view = df_view[
                df_view["Label"].str.lower().str.contains(s, na=False)
                | df_view["Value"].astype(str).str.lower().str.contains(s, na=False)
            ]
        if show_only_checked:
            df_view = df_view[df_view["Keep"] == True]
            if df_view.empty:
                st.info("Aucune ligne sélectionnée dans cet ensemble pour le moment.")
                return df

        editor_key = f"{key_prefix}_editor__{'checked' if show_only_checked else 'all'}"

        st.caption("La colonne Valeur est éditable (exemple : ajuster une course, une vitesse ou une charge).")
        edited = st.data_editor(
            df_view[["Keep", "Label", "Value", "Source", "_key"]],
            use_container_width=True,
            hide_index=True,
            height=900,
            column_config={
                "Keep": st.column_config.CheckboxColumn("Garder", default=False, width="small"),
                "Label": st.column_config.TextColumn("Variable", width="large"),
                "Value": st.column_config.TextColumn("Valeur (modifiable)", width="large"),
                "Source": st.column_config.TextColumn("Source", width="medium"),
                "_key": st.column_config.TextColumn("_key", width="small"),
            },
            key=editor_key,
        )

        df = _apply_editor_changes(df, edited)
        return df

    with tabs[0]:
        df1 = editor(df1, "ens1")
        st.session_state[df1_key] = df1

    if b2:
        with tabs[1]:
            df2 = editor(df2, "ens2")
            st.session_state[df2_key] = df2

    st.markdown("---")
    st.caption("Exporter la sélection finale en PDF et/ou Word (DOCX). Le document Word est entièrement modifiable.")
    if st.button("Générer les documents", type="primary", key="btn_generate_docs"):
        selected: List[PDFRow] = []

        def collect(df: pd.DataFrame, rows: List[PDFRow]):
            rows_map = {r.key: r for r in rows}
            checked = df[df["Keep"] == True]
            for _, r in checked.iterrows():
                obj = rows_map.get(str(r["_key"]))
                if obj:
                    obj.value = str(r["Value"])
                    obj.source = str(r["Source"])
                    selected.append(obj)

        collect(df1, rows1)
        collect(df2, rows2)

        if not selected:
            st.warning("Aucune ligne sélectionnée : veuillez cocher au moins une caractéristique.")
            return

        out_pdf = build_output_pdf(selected)
        out_docx = build_output_docx(selected)

        cdl1, cdl2 = st.columns(2)
        with cdl1:
            st.download_button(
                "💾 Télécharger le PDF",
                data=out_pdf,
                file_name="FICHE_TECHNIQUE.pdf",
                mime="application/pdf",
            )
        with cdl2:
            st.download_button(
                "📝 Télécharger le Word (DOCX)",
                data=out_docx,
                file_name="FICHE_TECHNIQUE.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
